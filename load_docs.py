import os
import io
import tempfile
from dotenv import load_dotenv
from langchain.document_loaders import TextLoader, PyPDFLoader
from langchain_community.vectorstores import PGVector
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as DocxDocument # Usar python-docx
from PIL import Image # Para manejar imágenes
import pytesseract # Para OCR

# Opcional: para usar un modelo multimodal de Google para captioning
from google.generativeai import GenerativeModel
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# Cargar variables de entorno
load_dotenv()

# --- Configuración ---
DB_HOST = os.getenv("DB_HOST", "localhost")
DB_PORT = os.getenv("DB_PORT", "5432")
DB_NAME = os.getenv("DB_NAME", "your_db_name")
DB_USER = os.getenv("DB_USER", "your_db_user")
DB_PASSWORD = os.getenv("DB_PASSWORD", "your_db_password")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY") # Necesaria si llamas directamente a la API para vision

COLLECTION_NAME = "knowledge_collection"
EMBEDDING_DIMENSION = 768 # for text-embedding-004

CONNECTION_STRING = f"postgresql://{DB_USER}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}"

# Directorio donde se encuentran tus documentos
DATA_DIRECTORY = "./docs" # <--- CAMBIA ESTO

# Configuración del Text Splitter
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

# Configuración de Tesseract OCR (cambia la ruta si es necesario)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' # Ejemplo Windows

# --- Inicializar Modelos ---
try:
    embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
    print("Modelo de Embedding de Texto inicializado.")
except Exception as e:
    print(f"Error al inicializar modelo de embedding: {e}")
    print("Asegúrate de que GOOGLE_API_KEY está configurada.")
    exit()

# Opcional: Inicializar modelo Vision si lo usas para captioning
try:
    vision_model = GenerativeModel("models/gemini-2.0-flash-thinking-exp-1219") # O gemini-1.5-pro
    print("Modelo Multimodal (Vision) inicializado.")
except Exception as e:
     print(f"Error al inicializar modelo Vision: {e}")
     print("Asegúrate de que GOOGLE_API_KEY está configurada y tienes acceso a 'gemini-pro-vision'.")
     vision_model = None # Continuar sin vision si falla


# --- Funciones de Procesamiento de Imagen ---

def extract_images_from_docx(docx_path):
    """Extrae texto e imágenes de un archivo .docx y devuelve una lista de partes."""
    try:
        doc = DocxDocument(docx_path)
        extracted_parts = []
        for paragraph in doc.paragraphs:
            extracted_parts.append({"type": "text", "content": paragraph.text})

        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                pass

        return extracted_parts
    except Exception as e:
        print(f"Error al extraer de {docx_path}: {e}")
        return []

def process_image(image_bytes):
    """
    Procesa bytes de imagen para extraer texto (OCR) y/o generar descripción (Captioning).
    Devuelve un texto combinado.
    """
    image_text = ""
    description = ""

    # 1. Intentar OCR
    try:
        img = Image.open(io.BytesIO(image_bytes))
        # Convierte a RGB si es necesario (pytesseract funciona mejor)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        ocr_text = pytesseract.image_to_string(img).strip()
        if ocr_text:
            image_text += f"Texto extraído de la imagen: {ocr_text}\n"
            print(f"  - OCR exitoso. Texto extraído: '{ocr_text[:100]}...'")
        else:
             print("  - OCR no extrajo texto.")

    except Exception as e:
        print(f"  - Error durante OCR: {e}")

    # 2. Opcional: Generar descripción con modelo multimodal
    # Esto requiere la API de Vision y puede ser costoso/lento.
    if vision_model:
        try:
            print("  - Generando descripción de la imagen con Vision Model...")
            image_part = {"mime_type": "image/jpeg", "data": image_bytes} # Ajusta el mime_type si no es jpeg
            prompt_parts = [image_part, "\nDescribe brevemente el contenido de esta imagen, incluyendo cualquier tabla o gráfico relevante."]
            response = vision_model.generate_content(
                prompt_parts,
                safety_settings={
                    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
                    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
                }
            )
            response.resolve() # Espera a que la respuesta esté disponible
            description = response.text.strip()
            if description:
                image_text += f"Descripción de la imagen: {description}\n"
                print(f"  - Descripción generada: '{description[:100]}...'")
            else:
                 print("  - Visión Model no generó descripción.")
        except Exception as e:
            print(f"  - Error al generar descripción con Vision Model: {e}")
            # Puedes añadir lógica de reintento o manejo específico de errores de la API
    else:
         print("  - Vision Model no inicializado o disponible. Omitiendo captioning.")


    return image_text.strip() # Devuelve el texto combinado de OCR y/o descripción


# --- Procesar Archivos ---

documents_for_embedding = []

# Asegurarse de que el directorio de datos existe
if not os.path.exists(DATA_DIRECTORY):
    print(f"Error: El directorio de datos '{DATA_DIRECTORY}' no existe.")
    print("Por favor, crea esta carpeta y coloca tus documentos dentro.")
    exit()

print(f"Procesando documentos desde '{DATA_DIRECTORY}'...")

# Recorrer los archivos en el directorio
for root, _, files in os.walk(DATA_DIRECTORY):
    for file_name in files:
        file_path = os.path.join(root, file_name)
        source_metadata = {"source": file_path}
        print(f"\nProcesando archivo: {file_name}")

        file_extension = os.path.splitext(file_name)[1].lower()

        # --- Lógica de carga y procesamiento según la extensión ---
        if file_extension == ".txt":
            try:
                loader = TextLoader(file_path)
                docs = loader.load()
                documents_for_embedding.extend(docs)
                print(f"  Cargado y añadido texto de {file_name}.")
            except Exception as e:
                print(f"  Error al cargar TXT {file_name}: {e}")

        elif file_extension == ".pdf":
             try:
                 loader = PyPDFLoader(file_path)
                 # PyPDFLoader ya maneja cada página como un Document.
                 # No extrae imágenes, pero el texto de cada página está listo para chunking.
                 docs = loader.load()
                 documents_for_embedding.extend(docs)
                 print(f"  Cargado y añadido texto de PDF {file_name} ({len(docs)} páginas).")
                 # NOTA: Para PDFs con imágenes, necesitarías librerías más avanzadas (ej. pdfplumber)
                 # para extraer imágenes y luego procesarlas con OCR/Vision.
             except Exception as e:
                 print(f"  Error al cargar PDF {file_name}: {e}")
                 print("  Asegúrate de tener 'pypdf' instalado.")

        elif file_extension == ".docx":
            try:
                print(f"  Extrayendo contenido de DOCX {file_name}...")
                extracted_parts = extract_images_from_docx(file_path)

                processed_text_with_images = ""
                for part in extracted_parts:
                    if part["type"] == "text":
                        processed_text_with_images += part["content"] + "\n"
                    elif part["type"] == "image":
                        # Procesar la imagen y añadir su texto/descripción al flujo
                        image_info_text = process_image(part["content"])
                        if image_info_text:
                             # Añadir un marcador claro para la información de la imagen
                             processed_text_with_images += f"\n--- Información de Imagen ---\n{image_info_text}\n--- Fin Información Imagen ---\n"

                # Crear un solo Document con todo el texto procesado
                if processed_text_with_images.strip():
                    # Mantener los metadatos originales
                    doc = Document(page_content=processed_text_with_images.strip(), metadata=source_metadata)
                    documents_for_embedding.append(doc)
                    print(f"  Procesado DOCX {file_name} (incluyendo texto e imágenes procesadas).")
                else:
                     print(f"  No se pudo extraer contenido procesable de DOCX {file_name}.")


            except Exception as e:
                print(f"  Error al procesar DOCX {file_name}: {e}")
                print("  Asegúrate de tener 'python-docx', 'pillow', 'pytesseract' y el motor Tesseract instalados.")


        # Añadir lógica para otros tipos de archivo si es necesario

print(f"\nTotal de documentos procesados antes de chunking: {len(documents_for_embedding)}")

# --- Dividir Documentos Procesados en Fragmentos ---
if not documents_for_embedding:
    print("No hay documentos procesados para generar embeddings.")
    exit()

print(f"Dividiendo documentos en fragmentos (chunk_size={CHUNK_SIZE}, chunk_overlap={CHUNK_OVERLAP})...")
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP
)
chunks = text_splitter.split_documents(documents_for_embedding)
print(f"Creados {len(chunks)} fragmentos.")

# --- Almacenar en PGVector ---
print(f"Conectando a la base de datos {DB_NAME} e insertando embeddings en tabla 'langchain_pg_embedding'...")

try:
    # Conectarse o crear la colección en PGVector
    # LangChain usará las tablas por defecto 'langchain_pg_collection' y 'langchain_pg_embedding'.
    vectorstore = PGVector(
        collection_name=COLLECTION_NAME,
        connection_string=CONNECTION_STRING,
        embedding_function=embeddings,
        # No pasar table_name si usas la estructura por defecto de LangChain
        use_jsonb=True
    )

    # Añadir los fragmentos al vector store. Esto genera los embeddings e inserta.
    print(f"Añadiendo {len(chunks)} fragmentos a la colección '{COLLECTION_NAME}'...")
    vectorstore.add_documents(chunks)

    print("Proceso de inserción completado.")

except Exception as e:
    print(f"Error al insertar embeddings en PGVector: {e}")
    print(f"Asegúrate de que la base de datos está corriendo y las tablas")
    print(f"'langchain_pg_collection' y 'langchain_pg_embedding' existen con la estructura esperada.")
    print("Verifica tus credenciales en el archivo .env")